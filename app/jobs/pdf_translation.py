from __future__ import annotations

import logging
import shutil
import zipfile
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Protocol, TypedDict

import requests
from flask import current_app


logger = logging.getLogger(__name__)


class PdfCompletedStatus(TypedDict):
    status: str
    filename: str
    stored_filename: str
    original_filename: str
    download_name: str
    message: str


class PdfFailedStatus(TypedDict):
    status: str
    error: str
    message: str


class MinerUData(TypedDict, total=False):
    task_id: str
    full_zip_url: str


class MinerUResult(TypedDict, total=False):
    code: int
    msg: str
    data: MinerUData


class PdfTaskStatusUpdater(Protocol):
    def completed(self, task_id: str, status: PdfCompletedStatus) -> None: ...
    def failed(self, task_id: str, status: PdfFailedStatus) -> None: ...


@dataclass(frozen=True, slots=True)
class PdfTranslationRequest:
    pdf_path: str
    original_filename: str
    unique_filename: str
    source_lang: str
    target_lang: str
    model: str
    enable_image_ocr: bool
    custom_translations: dict[str, str]
    user_id: int
    task_id: str
    output_path: str = ""
    register_history: bool = True


@dataclass(frozen=True, slots=True)
class PdfTranslationError(Exception):
    message: str

    def __str__(self) -> str:
        return self.message


def process_pdf_translation(request: PdfTranslationRequest, status_updater: PdfTaskStatusUpdater) -> bool:
    from app import create_app

    logger.info("start async PDF translation task: %s", request.task_id)
    app = create_app()
    with app.app_context():
        try:
            pdf_output_dir = _pdf_output_dir()
            task_work_dir = pdf_output_dir / f"{Path(request.unique_filename).stem}_work"
            task_work_dir.mkdir(parents=True, exist_ok=True)
            md_file = _extract_markdown(request, task_work_dir)
            docx_filename = _docx_filename(request)
            docx_path = Path(request.output_path) if request.output_path else pdf_output_dir / docx_filename
            docx_path.parent.mkdir(parents=True, exist_ok=True)
            if md_file:
                _translate_markdown(request, md_file, docx_path)
            else:
                _write_missing_markdown_doc(request, docx_path)
            if not docx_path.exists():
                raise PdfTranslationError(f"翻译后的文件不存在: {docx_path}")
            if request.register_history:
                _save_upload_record(request, docx_filename, pdf_output_dir, docx_path.stat().st_size)
            status_updater.completed(
                request.task_id,
                PdfCompletedStatus(
                    status="completed",
                    filename=docx_filename,
                    stored_filename=docx_filename,
                    original_filename=request.original_filename,
                    download_name=docx_filename,
                    message="翻译完成",
                ),
            )
            logger.info("PDF translation task completed: %s", request.task_id)
            return True
        except Exception as error:  # noqa: BROAD_EXCEPT_OK
            logger.error("PDF translation task failed: %s", error)
            logger.exception("error details")
            status_updater.failed(
                request.task_id,
                PdfFailedStatus(status="failed", error=str(error), message="翻译失败"),
            )
            raise


def _pdf_output_dir() -> Path:
    project_root = Path(__file__).resolve().parents[1]
    upload_folder = Path(current_app.config["UPLOAD_FOLDER"])
    if not upload_folder.is_absolute():
        upload_folder = project_root / upload_folder
    pdf_output_dir = upload_folder / "pdf_outputs"
    pdf_output_dir.mkdir(parents=True, exist_ok=True)
    return pdf_output_dir


def _extract_markdown(request: PdfTranslationRequest, task_work_dir: Path) -> Path | None:
    result = _extract_with_mineru(request)
    if result is None:
        result = _extract_with_local_processor(request)
    if not result:
        raise PdfTranslationError("所有PDF处理方法都失败了")
    if "code" in result and result["code"] != 0:
        raise PdfTranslationError(f"PDF处理失败: {result.get('msg', '未知错误')}")
    data = result.get("data")
    if not data or "task_id" not in data:
        raise PdfTranslationError("MinerU返回结果缺少task_id")
    if "full_zip_url" not in data:
        raise PdfTranslationError("MinerU返回结果缺少full_zip_url")
    mineru_task_id = data["task_id"]
    zip_path = task_work_dir / f"mineru_result_{mineru_task_id}.zip"
    _download_or_copy_zip(data["full_zip_url"], zip_path)
    with zipfile.ZipFile(zip_path, "r") as zip_ref:
        zip_ref.extractall(task_work_dir)
    return _find_markdown(task_work_dir, mineru_task_id)


def _extract_with_mineru(request: PdfTranslationRequest) -> MinerUResult | None:
    try:
        from app.function.image_ocr.ocr_api import MinerUAPI
        from app.function.image_ocr.oss_pdf_processor import OSSPDFProcessor

        result = OSSPDFProcessor().process_pdf_with_mineru(
            request.pdf_path,
            MinerUAPI(),
            bucket="fciai",
            region="cn-beijing",
            enable_ocr=request.enable_image_ocr,
        )
        if result and isinstance(result, dict) and result.get("code") == 0:
            return result
        logger.warning("OSS direct PDF processing failed, trying local processor")
        return None
    except Exception as error:  # noqa: BROAD_EXCEPT_OK
        logger.warning("OSS direct PDF processing failed: %s", error)
        return None


def _extract_with_local_processor(request: PdfTranslationRequest) -> MinerUResult:
    try:
        from app.function.local_pdf_processor import LocalPDFProcessor

        return LocalPDFProcessor().process_pdf(request.pdf_path)
    except Exception as error:  # noqa: BROAD_EXCEPT_OK
        logger.error("local PDF processor also failed: %s", error)
        raise PdfTranslationError("PDF处理失败，请检查文件格式") from error


def _download_or_copy_zip(zip_url: str, zip_path: Path) -> None:
    if zip_url.startswith("file://"):
        source_path = Path(zip_url[7:])
        if not source_path.exists():
            raise PdfTranslationError(f"源文件不存在: {source_path}")
        shutil.copy2(source_path, zip_path)
        return
    response = requests.get(zip_url, timeout=300, proxies={"http": None, "https": None})
    if response.status_code != 200:
        raise PdfTranslationError(f"下载ZIP文件失败，状态码: {response.status_code}")
    zip_path.write_bytes(response.content)


def _find_markdown(task_work_dir: Path, mineru_task_id: str) -> Path | None:
    markdown_files = [path for path in task_work_dir.rglob("*.md")]
    for path in markdown_files:
        if mineru_task_id in path.name:
            return path
    return markdown_files[0] if markdown_files else None


def _docx_filename(request: PdfTranslationRequest) -> str:
    original_base_name = Path(request.original_filename).stem
    return f"translated_{request.source_lang.lower()}_{request.target_lang.lower()}_{original_base_name}.docx"


def _translate_markdown(request: PdfTranslationRequest, md_file: Path, docx_path: Path) -> None:
    from app.utils.document_generator import translate_markdown_to_bilingual_doc

    content = md_file.read_text(encoding="utf-8")
    source_language = _normalize_pdf_language(request.source_lang, "en")
    target_language = _normalize_pdf_language(request.target_lang, "zh")
    ocr_results: list[dict[str, str | bool]] = []
    if request.enable_image_ocr:
        try:
            from app.function.image_ocr.ocr_controller import process_markdown_images_ocr_and_translate

            ocr_results = process_markdown_images_ocr_and_translate(
                markdown_content=content,
                markdown_dir=str(md_file.parent),
                target_language=target_language,
                source_language=source_language,
                provider_model=request.model,
            )
        except Exception as error:  # noqa: BROAD_EXCEPT_OK
            logger.error("PDF image OCR failed: %s", error)
            logger.exception("OCR error details")
    ok = translate_markdown_to_bilingual_doc(
        content,
        str(docx_path),
        source_language=source_language,
        target_language=target_language,
        image_base_dir=str(md_file.parent),
        custom_translations=request.custom_translations,
        image_ocr_results=ocr_results,
        provider_model=request.model,
    )
    if not ok:
        raise PdfTranslationError("翻译生成Word文档失败")


def _normalize_pdf_language(language: str, fallback: str) -> str:
    return {"EN": "en", "en": "en", "ZH": "zh", "zh": "zh", "JA": "ja", "ja": "ja"}.get(language, fallback)


def _write_missing_markdown_doc(request: PdfTranslationRequest, docx_path: Path) -> None:
    from docx import Document

    doc = Document()
    doc.add_heading("PDF处理结果", 1)
    doc.add_paragraph("未能从PDF中提取到文本内容，请检查原始PDF文件是否包含可提取的文本。")
    doc.add_paragraph(f"原始文件名: {request.original_filename}")
    doc.add_paragraph(f"处理时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.save(docx_path)


def _save_upload_record(request: PdfTranslationRequest, filename: str, file_path: Path, file_size: int) -> None:
    from app import db
    from app.models import UploadRecord

    record = UploadRecord(
        filename=filename,
        stored_filename=filename,
        file_path=str(file_path),
        user_id=request.user_id,
        file_size=file_size,
        status="completed",
    )
    db.session.add(record)
    db.session.commit()


__all__ = [
    "PdfCompletedStatus",
    "PdfFailedStatus",
    "PdfTaskStatusUpdater",
    "PdfTranslationError",
    "PdfTranslationRequest",
    "process_pdf_translation",
]
