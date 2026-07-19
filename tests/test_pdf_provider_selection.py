from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path

from docx import Document

from app.function.image_ocr.ocr_controller import _translate_text
from app.translation.types import ProviderError, ProviderName, ProviderRequest, ProviderResult
from app.utils.document_generator import translate_markdown_to_bilingual_doc


@dataclass(slots=True)
class RecordingProvider:
    provider_name: ProviderName
    prefix: str
    calls: list[ProviderRequest] = field(default_factory=list)

    @property
    def name(self) -> ProviderName:
        return self.provider_name

    def translate(self, request: ProviderRequest) -> ProviderResult:
        self.calls.append(request)
        return ProviderResult(
            text=f"{self.prefix}:{request.text}",
            provider=self.provider_name,
            model=self.provider_name,
        )


@dataclass(slots=True)
class FailingProvider:
    calls: list[ProviderRequest] = field(default_factory=list)

    @property
    def name(self) -> ProviderName:
        return "deepseek"

    def translate(self, request: ProviderRequest) -> ProviderResult:
        self.calls.append(request)
        raise ProviderError("deepseek", "provider_unavailable", "offline", retryable=True)


def _paragraph_shape(path: Path) -> list[str]:
    return [paragraph.style.name for paragraph in Document(path).paragraphs]


def test_markdown_provider_selection_keeps_document_structure(tmp_path: Path) -> None:
    markdown = "# Nutrition\n\nHuman milk supports development."
    qwen = RecordingProvider("qwen", "Q")
    deepseek = RecordingProvider("deepseek", "D")
    qwen_doc = tmp_path / "qwen.docx"
    deepseek_doc = tmp_path / "deepseek.docx"

    assert translate_markdown_to_bilingual_doc(markdown, str(qwen_doc), provider=qwen)
    assert translate_markdown_to_bilingual_doc(markdown, str(deepseek_doc), provider=deepseek)

    assert _paragraph_shape(qwen_doc) == _paragraph_shape(deepseek_doc)
    assert qwen.calls and deepseek.calls
    assert all(call.output_format == "plain" for call in qwen.calls + deepseek.calls)
    assert all(paragraph.text.startswith("Q:") for paragraph in Document(qwen_doc).paragraphs if paragraph.text.startswith("Q:"))
    assert all(paragraph.text.startswith("D:") for paragraph in Document(deepseek_doc).paragraphs if paragraph.text.startswith("D:"))


def test_provider_failure_does_not_fall_back(tmp_path: Path) -> None:
    provider = FailingProvider()
    output = tmp_path / "failed.docx"

    assert translate_markdown_to_bilingual_doc("One paragraph", str(output), provider=provider)

    assert len(provider.calls) == 1
    assert "One paragraph" in [paragraph.text for paragraph in Document(output).paragraphs]


def test_ocr_text_uses_injected_selected_provider() -> None:
    provider = RecordingProvider("deepseek", "OCR")

    translated = _translate_text("Milk", "zh", "en", provider_model="deepseek", provider=provider)

    assert translated == "OCR:Milk"
    assert provider.calls[0].field == "image OCR"
    assert provider.calls[0].output_format == "plain"
